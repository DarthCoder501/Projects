# Import necessary libraries
import re
from dateutil import parser
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import os
import pickle
from datetime import timedelta
from pypdf import PdfReader 
import docx

# Step 1: Extract text from syllabus (PDF or DOCX)
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        # Extract text from PDF
        reader = PdfReader(file_path)
        pages_num = len(reader.pages)

        # Collect all text from all pages
        extracted_text = "" 

        for i in range(pages_num):
            page = reader.pages[i]
            extracted_text += page.extract_text()  # Extract text and append it

        return extracted_text
    
    elif file_path.endswith('.docx'):
        # Extract text from DOCX
        doc = docx.Document(file_path)
        full_text = []

        # 1. Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)

        # 2. Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        # 3. Extract text from headers and footers
        section = doc.sections[0]  # Assumes first section for headers/footers
        header = section.header
        footer = section.footer
    
        for para in header.paragraphs:
            full_text.append(para.text)

        for para in footer.paragraphs:
            full_text.append(para.text)

        return '\n'.join(full_text)
    
    else:
        # Read plain text
        with open(file_path, 'r') as file:
            return file.read()

# Step 2: Extract dates and events from text
def extract_dates_events(text):
    date_event_pairs = []
    
    # Enhanced regex pattern to capture various date formats
    date_pattern = r"""
    (
        (?:                  # Start of the first capturing group
            (?:(?:\d{1,2})(?:st|nd|rd|th)?[ ]?(?:-|/)[ ]?(?:\d{1,2})(?:-|/)?(?:\d{2,4})?) |  # MM/DD/YYYY or DD/MM/YYYY with optional ordinals
            (?:\d{1,2}(?:st|nd|rd|th)?[ ]?(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[ ]?,?[ ]?(?:\d{2,4})?) |  # DD Month YYYY or DD Month, YYYY
            (?:[a-zA-Z]+[ ]\d{1,2}(?:st|nd|rd|th)?[ ]?,?[ ]?(?:\d{2,4})?)  # Month DD, YYYY or Month DD
        )
    )
    """
    
    # Compile the regex with verbose flag for readability
    date_pattern = re.compile(date_pattern, re.VERBOSE)
    
    # Find all matches for dates
    matches = re.finditer(date_pattern, text)

    last_pos = 0  # Track the position of the last match
    previous_date = None  # Track the previous date

    # Iterate over all date matches
    for match in matches:
        date_str = match.group()  # The date string found by regex
        
        # Parse the date into a standard datetime object
        try:
            parsed_date = parser.parse(date_str)  # Parse the date string into datetime object
        except Exception as e:
            # In case of parsing failure, skip the date
            continue

        # Extract the event description following the previous date
        if previous_date:
            event_description = text[last_pos:match.start()].strip()  # Text between the previous date and current date
            date_event_pairs.append((event_description, parsed_date))  # Store the event description and date

        # Update for the next iteration
        previous_date = parsed_date
        last_pos = match.end()

    # Append the final event after the last date match
    if previous_date:
        final_event = text[last_pos:].strip()  # Remaining text after the last date
        date_event_pairs.append((final_event, previous_date))

    return date_event_pairs

# Step 3: Create reminders for each event
def create_reminders(event_date):
    # Ensure event_date is a datetime object
    if isinstance(event_date, str):
        event_date = parser.parse(event_date)
    
    intervals = [1, 3, 7, 14, 30]  # 1 day, 3 days, 1 week, etc.
    reminders = [(event_date - timedelta(days=i)) for i in intervals]
    return reminders


# Step 4: Integrate with Google Calendar

# Google Calendar API authentication
SCOPES = ['https://www.googleapis.com/auth/calendar']

def authenticate_google_calendar():
    creds = None
    # Check if token.pickle exists and load credentials from it
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    
    # If no valid credentials, prompt the user to log in
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        
        # Save the credentials for the next run
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    
    # Build the Google Calendar API service
    service = build('calendar', 'v3', credentials=creds)
    return service


def add_to_google_calendar(event_name, event_date, reminders):
    service = authenticate_google_calendar()  # Authenticate and get the service
    
    # Define the event structure
    event = {
        'summary': event_name,  # Event title (e.g., "Read Ezekiel 1")
        'start': {
            'dateTime': event_date.isoformat(),  # Start date and time in ISO format
            'timeZone': 'America/New_York',   # Adjust the time zone
        },
        'end': {
            'dateTime': (event_date + timedelta(hours=1)).isoformat(),  # Assume 1-hour duration
            'timeZone': 'America/New_York',
        },
        'reminders': {
            'useDefault': False,  # Custom reminders
            'overrides': [
                {'method': 'popup', 'minutes': int((event_date - reminder).total_seconds() // 60)}
                for reminder in reminders  # Set reminders at the specified intervals
            ],
        },
    }
    
    # Insert the event into the primary calendar
    created_event = service.events().insert(calendarId='primary', body=event).execute()
    print(f'Event created: {created_event.get("htmlLink")}')

# Main function to run the process
def process_syllabus():
    # Get the file path from the user (prompted in the terminal)
    file_path = input("Please provide the full path to the syllabus file (PDF or DOCX): ")
    
    # Proceed with processing the file
    text = extract_text(file_path)
    date_event_pairs = extract_dates_events(text)
    
    for event_name, event_date in date_event_pairs:
        reminders = create_reminders(event_date)
        # Add to Google Calendar
        add_to_google_calendar(event_name, event_date, reminders)

# Run the main function
if __name__ == "__main__":
    process_syllabus()