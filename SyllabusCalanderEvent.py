import re  # Import the 're' module for working with regular expressions (used to search for patterns in text)
from dateutil import parser  # Import 'parser' from 'dateutil' to help convert text into date formats
from googleapiclient.discovery import build  # Import 'build' to create a service object for Google Calendar
from google.oauth2.credentials import Credentials  # Import 'Credentials' to manage Google account credentials
from google_auth_oauthlib.flow import InstalledAppFlow  # Import 'InstalledAppFlow' to manage Google OAuth2 authentication
from google.auth.transport.requests import Request  # Import 'Request' to handle HTTP requests during authentication
import os  # Import 'os' to interact with the operating system (e.g., checking if a file exists)
import pickle  # Import 'pickle' to save and load Google account credentials
from datetime import datetime, timedelta  # Import 'datetime' and 'timedelta' to work with dates and times
from pypdf import PdfReader  # Import 'PdfReader' to read text from PDF files
import docx  # Import 'docx' to read text from DOCX (Word) files
import logging  # Import 'logging' to record events and debug information

# Set up logging to record debug information and errors
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Function to extract text from a file (PDF or DOCX)
def extract_text(file_path):
    if file_path.endswith('.pdf'):  # Check if the file is a PDF
        reader = PdfReader(file_path)  # Create a PDF reader object
        extracted_text = ""  # Initialize an empty string to hold the extracted text
        for i in range(len(reader.pages)):  # Loop through each page in the PDF
            extracted_text += reader.pages[i].extract_text()  # Add the text from each page to 'extracted_text'
        return extracted_text  # Return all the extracted text
    
    elif file_path.endswith('.docx'):  # Check if the file is a DOCX (Word document)
        doc = docx.Document(file_path)  # Create a DOCX document object
        full_text = []  # Initialize an empty list to hold the extracted text
        for para in doc.paragraphs:  # Loop through each paragraph in the document
            full_text.append(para.text)  # Add the text of each paragraph to 'full_text'
        for table in doc.tables:  # Loop through each table in the document
            for row in table.rows:  # Loop through each row in each table
                for cell in row.cells:  # Loop through each cell in each row
                    full_text.append(cell.text)  # Add the text from each cell to 'full_text'
        return '\n'.join(full_text)  # Join all the text together into a single string
    
    else:  # If the file is a plain text file or another unsupported format
        with open(file_path, 'r') as file:  # Open the file in read mode
            return file.read()  # Return the content of the file as a string

# Function to identify the section of the syllabus that contains the course schedule
def identify_schedule_section(text):
    # List of keywords that might indicate the start of the course schedule section
    potential_keywords = ["COURSE SCHEDULE", "Timeline", "Important Dates", "Class Schedule", "Agenda"]
    for keyword in potential_keywords:  # Loop through each keyword
        # Search for the keyword in the text and capture everything until the end of the section
        match = re.search(rf'{keyword}(.*?)(?=\n\n|Final Exam:|End of Class)', text, re.DOTALL | re.IGNORECASE)
        if match:  # If a match is found
            return match.group(1).strip()  # Return the matching text (course schedule section)

    # If no clear match is found, log the potential sections and ask the user to select one
    logging.info("Potential schedule sections found:")
    sections = re.findall(r'(.*?(?:\n\d{1,2}\s\w+|\n\w+\s\d{1,2}))', text, re.DOTALL)
    for i, section in enumerate(sections):  # Loop through each potential section
        logging.info(f"Section {i+1}:")
        logging.info(section[:500])  # Log the first 500 characters of the section for context
        logging.info("\n" + "-"*40 + "\n")

    # Ask the user to select the correct section based on the logged information
    choice = int(input(f"Select the correct section (1-{len(sections)}), or 0 to abort: "))
    if 0 < choice <= len(sections):  # If the user selects a valid section number
        return sections[choice-1]  # Return the selected section
    else:  # If the user chooses to abort
        return None  # Return None to indicate no section was selected

# Function to extract dates and corresponding events from the text
def extract_dates_events(text):
    date_event_pairs = []  # Initialize an empty list to store pairs of events and dates
    lines = text.split('\n')  # Split the text into lines
    current_date = None  # Variable to store the current date

    for line in lines:  # Loop through each line in the text
        # Try to parse the date from the line
        try:
            date = parser.parse(line, fuzzy=True, default=datetime(2024, 1, 1))
            if date.year == 1:  # If the year is not specified, set it to 2024
                date = date.replace(year=2024)
            if date >= datetime(2024, 9, 1):  # Only consider dates from September 2024 onwards
                current_date = date  # Set the current date to the parsed date
                continue  # Skip to the next line after setting the date
        except ValueError:
            pass  # If the line does not contain a date, continue to the next line

        # If a valid date is found and the line is not empty, treat it as an event
        if current_date and line.strip():
            date_event_pairs.append((line.strip(), current_date))  # Add the event and date to the list
            logging.debug(f"Extracted: Event '{line.strip()}' on {current_date}")  # Log the extracted event and date
    
    return date_event_pairs  # Return the list of extracted date-event pairs

# Function to create reminder times for each event based on its date
def create_reminders(event_date):
    intervals = [1, 3, 7, 14, 30]  # Days before the event to create reminders (1 day, 3 days, etc.)
    reminders = [(event_date - timedelta(days=i)) for i in intervals]  # Calculate reminder dates based on these intervals
    return reminders  # Return the list of reminder dates

# Function to authenticate with Google Calendar
SCOPES = ['https://www.googleapis.com/auth/calendar']

def authenticate_google_calendar():
    creds = None  # Initialize the credentials variable
    if os.path.exists('token.pickle'):  # Check if a credentials file already exists
        with open('token.pickle', 'rb') as token:  # If it exists, open it
            creds = pickle.load(token)  # Load the credentials from the file

    if not creds or not creds.valid:  # If there are no valid credentials
        if creds and creds.expired and creds.refresh_token:  # If the credentials have expired but can be refreshed
            creds.refresh(Request())  # Refresh the credentials
        else:  # If not, the user needs to log in again
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)  # Start a new login flow
            creds = flow.run_local_server(port=0)  # Open a local server to complete the login

        with open('token.pickle', 'wb') as token:  # Save the new credentials to a file
            pickle.dump(creds, token)  # Store the credentials for future use

    service = build('calendar', 'v3', credentials=creds)  # Build the Google Calendar service using the credentials
    return service  # Return the service object

# Function to add events to Google Calendar
def add_to_google_calendar(event_name, event_date, reminders):
    service = authenticate_google_calendar()  # Authenticate and get the Google Calendar service
    event = {
        'summary': event_name,  # The name of the event
        'start': {'dateTime': event_date.isoformat(), 'timeZone': 'America/New_York'},  # Start time and time zone
        'end': {'dateTime': (event_date + timedelta(hours=1)).isoformat(), 'timeZone': 'America/New_York'},  # End time (1 hour later)
        'reminders': {
            'useDefault': False,  # Don't use default reminders
            'overrides': [{'method': 'popup', 'minutes': int((event_date - reminder).total_seconds() // 60)} for reminder in reminders],  # Set custom reminders based on the intervals
        },
    }
    try:
        created_event = service.events().insert(calendarId='primary', body=event).execute()  # Insert the event into the user's primary calendar
        logging.info(f'Event created: {created_event.get("htmlLink")}')  # Log the link to the created event
    except Exception as e:  # Catch any errors that occur during event creation
        logging.error(f"Failed to create event: {event_name} on {event_date}. Error: {str(e)}")  # Log the error message

# Main function to process the syllabus and add events to Google Calendar
def process_syllabus():
    file_path = input("Please provide the full path to the syllabus file (PDF or DOCX): ")  # Ask the user for the file path
    text = extract_text(file_path)  # Extract text from the provided file
    course_schedule_text = identify_schedule_section(text)  # Identify the section of the text that contains the course schedule
    
    if course_schedule_text:  # If the course schedule section was found
        logging.info(f"Identified Schedule Section:\n{course_schedule_text[:500]}...")  # Log the start of the schedule section
        
        date_event_pairs = extract_dates_events(course_schedule_text)  # Extract dates and events from the schedule
        logging.info(f"Extracted {len(date_event_pairs)} date-event pairs")  # Log the number of extracted pairs
        
        for event_name, event_date in date_event_pairs:  # Loop through each event and date
            reminders = create_reminders(event_date)  # Create reminders for each event
            logging.info(f"Adding event: {event_name} on {event_date}")  # Log the event being added
            add_to_google_calendar(event_name, event_date, reminders)  # Add the event to Google Calendar
    else:
        logging.warning("Could not identify the course schedule section.")  # If no schedule section was found, log a warning

# If this script is run directly, start the process
if __name__ == "__main__":
    process_syllabus()
